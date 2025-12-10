import io
import json
import asyncio
import logging
from typing import List
from io import BytesIO

from fastapi import APIRouter, File, UploadFile, HTTPException
from PIL import Image
import pytesseract
import pdfplumber
from docx import Document
from sentence_transformers import SentenceTransformer, util
from keybert import KeyBERT

from models.chat import ChatResponse
from models.cv import CVAnalysisResponse
from services.chat_service import CVChatService
from services.embedding_service import LocalEmbeddingService
from services.vector_store import FAISSVectorStore
from services.gpt_service import GPTService
from db.job_repository import get_complete_job_details
from config.settings import JSON_FILE_PATH

logger = logging.getLogger(__name__)

router = APIRouter()

# Initialize services (will be injected from main app)
embedding_service: LocalEmbeddingService = None
vector_store: FAISSVectorStore = None
gpt_service: GPTService = None
cv_chat_service: CVChatService = None
cv_processor = None

# Initialize model for legacy resume matching
model = SentenceTransformer("all-MiniLM-L6-v2")

# Load KeyBERT model once
kw_model = KeyBERT("models/all-MiniLM-L6-v2")


def init_services(embed_svc, vec_store, gpt_svc, cv_chat_svc, cv_proc):
    """Initialize services from main app"""
    global embedding_service, vector_store, gpt_service, cv_chat_service, cv_processor
    embedding_service = embed_svc
    vector_store = vec_store
    gpt_service = gpt_svc
    cv_chat_service = cv_chat_svc
    cv_processor = cv_proc


def _generate_cv_recommendations(profile) -> List[str]:
    """Generate recommendations based on CV analysis"""
    recommendations = []

    # Skills recommendations
    if len(profile.skills) < 5:
        recommendations.append("Add more technical skills to improve job matching")

    # Contact information
    if not profile.email:
        recommendations.append("Add contact email for better profile completeness")

    if not profile.phone:
        recommendations.append("Include phone number in your CV")

    # Experience recommendations
    if len(profile.experience) == 0:
        recommendations.append("Add work experience details for better job matching")
    elif len(profile.experience) < 2:
        recommendations.append("Include more work experience entries if available")

    # Education recommendations
    if len(profile.education) == 0:
        recommendations.append("Add education background to strengthen your profile")

    # Summary recommendations
    if not profile.summary:
        recommendations.append("Add a professional summary to highlight your strengths")

    # Confidence-based recommendations
    if profile.confidence_score < 0.5:
        recommendations.append("Consider adding more detailed information to improve CV quality")
    elif profile.confidence_score < 0.7:
        recommendations.append("Good CV structure! Add a few more details for optimal results")

    # Skills gap analysis
    high_demand_skills = [
        "Python", "JavaScript", "React", "Node.js", "AWS", "Docker",
        "Kubernetes", "SQL", "MongoDB", "Git", "CI/CD"
    ]

    missing_skills = [skill for skill in high_demand_skills
                     if skill.lower() not in [s.lower() for s in profile.skills]]

    if missing_skills:
        recommendations.append(f"Consider learning in-demand skills: {', '.join(missing_skills[:3])}")

    return recommendations[:5]  # Limit to top 5 recommendations


@router.post("/upload_cv", response_model=CVAnalysisResponse)
async def upload_cv_enhanced(cv_file: UploadFile = File(...)) -> CVAnalysisResponse:
    """Enhanced CV upload with complete analysis and job matching"""
    start_time = asyncio.get_event_loop().time()

    try:
        if not cv_file.filename:
            raise HTTPException(status_code=400, detail="No file provided")

        # Validate file type
        allowed_extensions = {'.pdf', '.doc', '.docx', '.png', '.jpg', '.jpeg'}
        file_ext = '.' + cv_file.filename.lower().split('.')[-1]

        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )

        logger.info(f"Processing CV upload: {cv_file.filename}")

        # Read file content
        file_content = await cv_file.read()

        # Process CV using enhanced processor
        cv_profile = await cv_processor.process_cv(file_content, cv_file.filename)

        # Get job search text
        search_text = cv_processor.get_job_search_text(cv_profile)

        jobs_found = []
        total_jobs = 0

        # Perform job search if we have meaningful skills
        if cv_profile.skills and len(cv_profile.skills) >= 2:
            try:
                # Generate embedding for combined profile text
                profile_embedding = await embedding_service.get_embedding(search_text)

                # Search similar jobs
                similar_jobs = await vector_store.search_similar_jobs(
                    profile_embedding,
                    top_k=30
                )

                if similar_jobs:
                    # Re-rank jobs using GPT with extracted skills
                    ranked_jobs = await gpt_service.rerank_jobs(cv_profile.skills, similar_jobs)

                    # Get complete job details for top matches
                    job_ids = [job["ncspjobid"] for job in ranked_jobs[:10]]
                    complete_jobs = await get_complete_job_details(job_ids)

                    # Format job results
                    for job_data in ranked_jobs[:10]:
                        complete_job = next(
                            (j for j in complete_jobs if j["ncspjobid"] == job_data["ncspjobid"]),
                            {}
                        )

                        if complete_job:
                            jobs_found.append({
                                "ncspjobid": job_data["ncspjobid"],
                                "title": job_data["title"],
                                "organization_name": complete_job.get("organization_name", ""),
                                "match_percentage": job_data["match_percentage"],
                                "statename": complete_job.get("statename", ""),
                                "districtname": complete_job.get("districtname", ""),
                                "avewage": complete_job.get("avewage", 0),
                                "aveexp": complete_job.get("aveexp", 0),
                                "keywords": complete_job.get("keywords", ""),
                                "functionalrolename": complete_job.get("functionalrolename", ""),
                                "industryname": complete_job.get("industryname", ""),
                                "skills_matched": job_data.get("keywords_matched", []),
                                "similarity_score": job_data.get("similarity_used", 0)
                            })

                    total_jobs = len(ranked_jobs)

            except Exception as job_search_error:
                logger.error(f"Job search failed during CV processing: {job_search_error}")
                # Continue without job results

        # Generate processing recommendations
        recommendations = _generate_cv_recommendations(cv_profile)

        processing_time_ms = int((asyncio.get_event_loop().time() - start_time) * 1000)

        # Success response
        success_message = f"Successfully processed your CV! "

        if jobs_found:
            success_message += f"Found {len(jobs_found)} matching jobs with {len(cv_profile.skills)} extracted skills."
        else:
            success_message += f"Extracted {len(cv_profile.skills)} skills. Try refining your CV for better job matches."

        logger.info(f"CV processing completed: {cv_profile.confidence_score} confidence, "
                   f"{len(jobs_found)} jobs, {processing_time_ms}ms")

        return CVAnalysisResponse(
            success=True,
            message=success_message,
            profile=cv_processor.to_dict(cv_profile),
            jobs=jobs_found,
            total_jobs_found=total_jobs,
            processing_time_ms=processing_time_ms,
            confidence_score=cv_profile.confidence_score,
            recommendations=recommendations
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"CV upload processing failed: {e}")
        processing_time_ms = int((asyncio.get_event_loop().time() - start_time) * 1000)

        return CVAnalysisResponse(
            success=False,
            message=f"Failed to process CV: {str(e)}",
            processing_time_ms=processing_time_ms
        )


@router.post("/analyze_cv", response_model=CVAnalysisResponse)
async def analyze_cv_only(cv_file: UploadFile = File(...)) -> CVAnalysisResponse:
    """Analyze CV structure and extract data without job matching"""
    start_time = asyncio.get_event_loop().time()

    try:
        if not cv_file.filename:
            raise HTTPException(status_code=400, detail="No file provided")

        # Validate file type
        allowed_extensions = {'.pdf', '.doc', '.docx', '.png', '.jpg', '.jpeg'}
        file_ext = '.' + cv_file.filename.lower().split('.')[-1]

        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )

        logger.info(f"Analyzing CV: {cv_file.filename}")

        # Read and process CV
        file_content = await cv_file.read()
        cv_profile = await cv_processor.process_cv(file_content, cv_file.filename)

        # Generate recommendations
        recommendations = _generate_cv_recommendations(cv_profile)

        processing_time_ms = int((asyncio.get_event_loop().time() - start_time) * 1000)

        analysis_message = f"CV Analysis Complete! Extracted {len(cv_profile.skills)} skills, "
        analysis_message += f"{len(cv_profile.experience)} experience entries, "
        analysis_message += f"confidence score: {cv_profile.confidence_score:.1%}"

        return CVAnalysisResponse(
            success=True,
            message=analysis_message,
            profile=cv_processor.to_dict(cv_profile),
            processing_time_ms=processing_time_ms,
            confidence_score=cv_profile.confidence_score,
            recommendations=recommendations
        )

    except Exception as e:
        logger.error(f"CV analysis failed: {e}")
        processing_time_ms = int((asyncio.get_event_loop().time() - start_time) * 1000)

        return CVAnalysisResponse(
            success=False,
            message=f"Analysis failed: {str(e)}",
            processing_time_ms=processing_time_ms
        )


@router.post("/upload_cv_chat", response_model=ChatResponse)
async def upload_cv_for_chat(cv_file: UploadFile = File(...)) -> ChatResponse:
    """Upload CV and get chat-style response with job matches"""
    try:
        if not cv_file.filename:
            raise HTTPException(status_code=400, detail="No file provided")

        # Process CV using the enhanced processor
        file_content = await cv_file.read()
        cv_profile = await cv_processor.process_cv(file_content, cv_file.filename)

        # Generate chat response with job matches using CV service
        chat_response = await cv_chat_service.handle_cv_upload_chat(cv_profile)

        return chat_response

    except Exception as e:
        logger.error(f"CV chat upload failed: {e}")
        return ChatResponse(
            response="I had trouble processing your CV. Let's build your profile by chatting about your skills!",
            message_type="text",
            chat_phase="profile_building"
        )


# =========================================================================
# LEGACY RESUME UPLOAD ENDPOINT
# =========================================================================

@router.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    job_matches = []  # initialize here

    content = await file.read()
    cv_text, table_output = process_uploaded_file(content, file.filename)
    cv_keywords_text = extract_cv_keywords(cv_text, table_rows=table_output, top_n=30)

    if cv_keywords_text and cv_keywords_text.strip():   # only if CV text exists
        # Read job JSON
        jobs_data = read_json()
        if jobs_data["status"] == "error":
            return jobs_data

        job_matches = match_jobs(cv_keywords_text, jobs_data["data"]["content"])
        print(f" Found {len(job_matches)} job matches for CV: {file.filename}")
    return {
        "filename": file.filename,
        "size": len(content),
        "matches": job_matches  # return matches in response
    }


def process_uploaded_file(file_bytes: bytes, filename: str):
    text_output = ""
    table_output = []
    file_ext = filename.lower().split('.')[-1]

    try:
        if file_ext in ['png', 'jpg', 'jpeg']:
            text_output = perform_ocr_on_image(file_bytes)

        elif file_ext == 'pdf':
            text_output, table_output = extract_text_from_pdf(file_bytes)
            print("PDF file processing is not implemented yet.")

        elif file_ext == 'docx':
            text_output = extract_text_from_docx(file_bytes)

        else:
            print("unsupported format")

    except Exception as e:
        print("Error processing file:", str(e))
    return text_output, table_output


def perform_ocr_on_image(image_bytes):
    img = Image.open(io.BytesIO(image_bytes))
    text = pytesseract.image_to_string(img)
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    extracted_text = []

    # Extract paragraphs (handles text + bullets/numbers)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            extracted_text.append(text)

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_data.append(cell_text)
            if row_data:
                extracted_text.append(" | ".join(row_data))

    return "\n".join(extracted_text)


def extract_text_from_pdf(file_bytes):
    text_output = ""
    table_output = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        print(f"PDF opened successfully, total pages: {len(pdf.pages)}")

        for i, page in enumerate(pdf.pages):
            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_data = [cell.strip() for cell in row if cell and cell.strip()]
                    if row_data:
                        table_line = " | ".join(row_data)
                        text_output += f"\n--- Page {i+1} Table ---\n{table_line}"
                        table_output.append({
                            "page_number": i + 1,
                            "row_text": table_line
                        })

            # Extract text
            text = page.extract_text()
            if text:
                text_output += f"\n--- Page {i+1} ---\n{text.strip()}"
            else:
                pil_image = page.to_image(resolution=300).original
                ocr_text = pytesseract.image_to_string(pil_image)
                if ocr_text.strip():
                    text_output += f"\n--- OCR Page {i+1} ---\n{ocr_text.strip()}"
        print(f"\n Finished PDF processing. Total extracted text length: {len(text_output)} chars")
        print(f"Total table rows extracted: {len(table_output)}")
    return text_output, table_output


def read_json():
    try:
        with open(JSON_FILE_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
        return {"status": "success", "data": data}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def match_jobs(cv_text: str, jobs: list, top_k: int = 100):
    # Only use skills (and optionally role) for similarity
    job_texts = []
    for job in jobs:
        skills = job.get("key_skills", "") or ""
        role_text = job.get("role", "") or ""

        # Search only based on skills (+ optional role if you want)
        job_doc = f"Skills: {skills}. Role: {role_text}"
        job_texts.append(job_doc)

    # Encode CV and Jobs
    cv_embedding = model.encode(cv_text, convert_to_tensor=True)
    job_embeddings = model.encode(job_texts, convert_to_tensor=True)

    # Similarity search
    similarities = util.cos_sim(cv_embedding, job_embeddings)[0]
    top_results = similarities.topk(k=min(top_k, len(job_texts)))

    # Return ALL fields in output, not just skills
    matches = []
    for score, idx in zip(top_results.values, top_results.indices):
        matches.append({
            **jobs[idx],         # includes all job fields (title, salary, etc.)
            "score": float(score)  # similarity score
        })
    return matches


def extract_cv_keywords(cv_text: str, table_rows: list = None, top_n: int = 30) -> str:
    # Combine paragraph text and table rows
    combined_text = cv_text
    if table_rows:
        combined_text += "\n" + "\n".join(
            [" ".join(map(str, row)) for row in table_rows]
        )

    # Extract top keywords using KeyBERT
    keywords_with_scores = kw_model.extract_keywords(
        combined_text,
        top_n=top_n,
        stop_words="english"
    )

    # Convert to a single string
    keywords_text = " ".join([kw for kw, score in keywords_with_scores])

    return keywords_text
